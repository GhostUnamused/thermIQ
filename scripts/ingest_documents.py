"""
Local one-time ingestion script: document -> chunks -> Jina embeddings -> Qdrant + Firestore.

Supported formats (detected from the file extension): .pdf, .docx, .xlsx, .csv, .txt

Usage:
  python scripts/ingest_documents.py <file_path> <doc_type> <source_doc_name> <source_url> [client_name]

  client_name: optional. Pass the plant/client slug (e.g. "ntpc") for plant-specific docs.
               Omit (or pass "") for benchmark/regulatory docs (CEA specs etc.).

Optional env vars (set by the Drive ingest worker):
  THERMIQ_DRIVE_FILE_ID / THERMIQ_ORIGIN_FOLDER_ID — provenance written onto the
      Firestore document record (enables Drive folder re-sync diffing).
  THERMIQ_SKIP_RELEVANCE=1 — bypass the Gemini relevance gate.
  GEMINI_API_KEY — enables the relevance gate (client docs only; skipped if unset).

Exit codes: 0 ok · 1 error · 3 rejected by relevance screening.
"""
import json
import os
import sys
import time
import uuid
from datetime import datetime

import requests
from dotenv import load_dotenv
from pypdf import PdfReader
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, PointStruct, VectorParams
import firebase_admin
from firebase_admin import credentials, firestore

load_dotenv()

JINA_API_KEY = os.environ.get("JINA_API_KEY")
QDRANT_URL = os.environ.get("QDRANT_URL")
QDRANT_API_KEY = os.environ.get("QDRANT_API_KEY")

COLLECTION_NAME = "thermiq_chunks"
CHUNK_WORDS = 400
CHUNK_OVERLAP = 50

EQUIPMENT_KEYWORDS = {
    "Boiler": [
        "boiler",
        "furnace",
        "super heater",
        "superheater",
        "economiser",
        "air preheater",
        "steam drum",
        "burner",
    ],
    "Turbine": [
        "turbine",
        "hp turbine",
        "lp turbine",
        "ip turbine",
        "blade",
        "governor",
        "nozzle",
        "rotor",
    ],
    "Generator": [
        "generator",
        "stator",
        "rotor winding",
        "exciter",
        "avr",
        "alternator",
    ],
    "BFP": ["boiler feed pump", "bfp", "feed pump", "seal water", "impeller"],
    "Condenser": ["condenser", "cooling water", "hotwell", "vacuum", "ejector"],
    "Cooling Tower": [
        "cooling tower",
        "ct fill",
        "drift eliminator",
        "basin",
        "blowdown",
    ],
}


def extract_pages(pdf_path):
    reader = PdfReader(pdf_path)
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if len(text.strip()) < 50:
            continue
        pages.append((i, text))
    return pages


def extract_document(file_path):
    """Format-aware extraction. Returns (pages, total_pages) where pages is a
    list of (page_number, text). Non-paginated formats return one pseudo-page
    per logical unit (docx: whole doc; xlsx: per sheet; csv/txt: whole file)."""
    ext = os.path.splitext(file_path)[1].lower().lstrip(".")

    if ext == "pdf":
        total = len(PdfReader(file_path).pages)
        return extract_pages(file_path), total

    if ext == "docx":
        from docx import Document  # python-docx

        doc = Document(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return ([(1, text)] if text.strip() else []), 1

    if ext == "xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        pages = []
        for idx, ws in enumerate(wb.worksheets, start=1):
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(v) for v in row if v is not None and str(v).strip()]
                if cells:
                    rows.append(", ".join(cells))
            if rows:
                pages.append((idx, f"[Sheet: {ws.title}]\n" + "\n".join(rows)))
        return pages, len(wb.worksheets)

    if ext in ("csv", "txt"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return ([(1, text)] if text.strip() else []), 1

    raise ValueError(f"Unsupported file type '.{ext}' — supported: pdf, docx, xlsx, csv, txt")


def relevance_check(text, doc_name):
    """Gemini gatekeeper: is this document plausibly thermal-power-plant material?
    Returns (relevant, reason). Fails OPEN on any API problem."""
    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not api_key or os.environ.get("THERMIQ_SKIP_RELEVANCE") == "1":
        return True, "screening skipped"
    try:
        resp = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}",
            json={
                "contents": [{"parts": [{"text": (
                    "You are a gatekeeper for a thermal power plant knowledge base (SOPs, equipment "
                    "manuals, CEA/CERC regulations, tariff petitions, inspection & maintenance records, "
                    "outage reports, plant engineering).\n"
                    "Decide if the following document belongs in that knowledge base. Industrial/"
                    "engineering/energy-sector documents count as relevant; generic unrelated content "
                    "(event handbooks, marketing, fiction, recipes, unrelated coursework) does not.\n"
                    'Respond with ONLY a JSON object: {"relevant": true|false, "reason": "<one short sentence>"}\n\n'
                    f"Document name: {doc_name}\n\nDocument text sample:\n{text[:6000]}"
                )}]}],
                "generationConfig": {"temperature": 0, "maxOutputTokens": 100, "responseMimeType": "application/json"},
            },
            timeout=30,
        )
        resp.raise_for_status()
        raw = resp.json()["candidates"][0]["content"]["parts"][0]["text"]
        verdict = json.loads(raw)
        return bool(verdict.get("relevant", True)), verdict.get("reason", "")
    except Exception as e:  # noqa: BLE001 — screening must never block on outages
        print(f"Relevance screening skipped (non-fatal): {e}")
        return True, "screening unavailable"


def chunk_pages(pages, doc_slug):
    chunks = []
    chunk_index = 0

    full_words = []
    for page_number, text in pages:
        for word in text.split():
            full_words.append((word, page_number))

    i = 0
    while i < len(full_words):
        window = full_words[i : i + CHUNK_WORDS]
        if not window:
            break
        words = [w for w, _ in window]
        page_number = window[0][1]
        chunk_text = " ".join(words)
        chunk_id = f"{doc_slug}_chunk_{chunk_index:04d}"
        chunks.append(
            {
                "chunk_id": chunk_id,
                "chunk_index": chunk_index,
                "page_number": page_number,
                "text": chunk_text,
            }
        )
        chunk_index += 1
        i += CHUNK_WORDS - CHUNK_OVERLAP

    return chunks


def extract_equipment_tags(text):
    lowered = text.lower()
    matched = []
    for equipment, keywords in EQUIPMENT_KEYWORDS.items():
        if any(keyword in lowered for keyword in keywords):
            matched.append(equipment)
    return matched


def embed_batch(texts):
    for attempt in range(6):
        response = requests.post(
            "https://api.jina.ai/v1/embeddings",
            headers={
                "Authorization": f"Bearer {JINA_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": "jina-embeddings-v3",
                "input": texts,
                "task": "retrieval.passage",
            },
        )
        if response.status_code == 429 and attempt < 5:
            wait = 2 ** (attempt + 2)
            print(f"  Jina 429 rate limit, retrying in {wait}s ...")
            time.sleep(wait)
            continue
        response.raise_for_status()
        break
    data = response.json()["data"]
    return [item["embedding"] for item in data]


def get_firestore_client():
    try:
        app = firebase_admin.get_app("thermiq")
    except ValueError:
        cred = credentials.Certificate(
            {
                "type": "service_account",
                "project_id": os.environ["FIREBASE_PROJECT_ID"],
                "private_key": os.environ["FIREBASE_PRIVATE_KEY"].replace(
                    "\\n", "\n"
                ),
                "client_email": os.environ["FIREBASE_CLIENT_EMAIL"],
                "token_uri": "https://oauth2.googleapis.com/token",
            }
        )
        app = firebase_admin.initialize_app(cred, name="thermiq")
    return firestore.client(app=app)


def main():
    if len(sys.argv) < 5:
        print(
            "Usage: python scripts/ingest_documents.py <pdf_path> <doc_type> "
            "<source_doc_name> <source_url> [client_name]"
        )
        print("  client_name: optional, e.g. 'ntpc' for plant-specific docs (default: benchmark)")
        sys.exit(1)

    pdf_path, doc_type, source_doc_name, source_url = sys.argv[1:5]
    doc_client = (sys.argv[5] if len(sys.argv) > 5 else "").strip().lower()
    # Derive source_type from whether a client name was provided
    source_type = "client" if doc_client else "benchmark"
    client_name = doc_client  # "" for benchmark docs
    doc_slug = source_doc_name.lower().replace(" ", "_")

    start_time = time.time()

    print(f"Extracting text from {pdf_path} ...")
    pages, total_pages = extract_document(pdf_path)
    print(f"Extracted {len(pages)} usable pages/sections out of {total_pages} total.")
    if not pages:
        print("No extractable text — aborting.", file=sys.stderr)
        sys.exit(1)

    # Relevance gate (client docs only — benchmarks are curated by hand).
    if source_type == "client":
        sample = "\n".join(t for _, t in pages[:5])
        relevant, reason = relevance_check(sample, source_doc_name)
        if not relevant:
            print(f"REJECTED by relevance screening: {reason}", file=sys.stderr)
            sys.exit(3)

    print("Chunking text ...")
    chunks = chunk_pages(pages, doc_slug)
    print(f"Created {len(chunks)} chunks.")

    print("Tagging equipment per chunk ...")
    for chunk in chunks:
        chunk["equipment_tags"] = extract_equipment_tags(chunk["text"])

    print("Embedding chunks in batches of 8 ...")
    embeddings = []
    for i in range(0, len(chunks), 8):
        batch = chunks[i : i + 8]
        batch_texts = [c["text"] for c in batch]
        batch_embeddings = embed_batch(batch_texts)
        embeddings.extend(batch_embeddings)
        print(f"  Embedded batch {i // 8 + 1} ({len(batch)} chunks).")
        time.sleep(1)

    print("Connecting to Qdrant ...")
    client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=120)

    existing_collections = [c.name for c in client.get_collections().collections]
    if COLLECTION_NAME not in existing_collections:
        print(f"Creating collection '{COLLECTION_NAME}' ...")
        client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(size=1024, distance=Distance.COSINE),
        )
        # Qdrant Cloud requires an explicit payload index before a field can be used
        # in a filter — without these, every source_type/client_name filtered query
        # (gap detection, query.js benchmark/client split) fails with a 400.
        client.create_payload_index(COLLECTION_NAME, field_name="source_type", field_schema="keyword")
        client.create_payload_index(COLLECTION_NAME, field_name="client_name", field_schema="keyword")

    print("Building points and saving chunk fallback files ...")
    os.makedirs("data/chunks", exist_ok=True)
    points = []
    ingested_at = datetime.utcnow().isoformat() + "Z"
    for chunk, embedding in zip(chunks, embeddings):
        payload = {
            "chunk_id": chunk["chunk_id"],
            "source_doc": source_doc_name,
            "source_url": source_url,
            "doc_type": doc_type,
            "client": doc_client,  # "" for benchmark docs; "ntpc" etc. for plant-specific
            "source_type": source_type,   # "benchmark" or "client" — used by gap detection filter
            "client_name": client_name,   # "" for benchmark; plant slug for client docs
            "equipment_tags": chunk["equipment_tags"],
            "section": "",
            "page_number": chunk["page_number"],
            "text": chunk["text"],
            "chunk_index": chunk["chunk_index"],
            "ingested_at": ingested_at,
        }
        points.append(
            PointStruct(id=str(uuid.uuid4()), vector=embedding, payload=payload)
        )

        with open(f"data/chunks/{chunk['chunk_id']}.json", "w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2)

    print("Upserting to Qdrant in batches of 20 ...")
    upsert_count = 0
    for i in range(0, len(points), 20):
        batch = points[i : i + 20]
        client.upsert(collection_name=COLLECTION_NAME, points=batch)
        upsert_count += len(batch)
        print(f"  Upserted {upsert_count}/{len(points)} points.")

    print("Updating Firestore system_meta + documents ...")
    db = get_firestore_client()

    # 1 — Update aggregate counters on system_meta/config (existing behaviour)
    config_ref = db.collection("system_meta").document("config")
    config_ref.set(
        {
            "documents_ingested": firestore.ArrayUnion([source_doc_name]),
            "total_chunks_indexed": firestore.Increment(len(chunks)),
            "last_ingestion_at": ingested_at,
        },
        merge=True,
    )

    # 2 — Write a per-document record to the `documents` collection.
    #     This is what api/list_documents.js reads to populate the Documents page.
    #     Without this, locally-ingested docs are live in Qdrant but invisible in the UI.
    doc_id = f"{source_type}_{int(time.time() * 1000)}"
    doc_record = {
        "doc_name": source_doc_name,
        "doc_type": doc_type,
        "client": doc_client,
        "source_url": source_url,
        "source_type": source_type,
        "client_name": client_name,
        "chunks_indexed": len(chunks),
        "pages_parsed": total_pages,
        "ingested_at": ingested_at,
    }
    # Drive provenance (set by scripts/ingest_from_drive.py) — enables folder re-sync.
    if os.environ.get("THERMIQ_DRIVE_FILE_ID"):
        doc_record["drive_file_id"] = os.environ["THERMIQ_DRIVE_FILE_ID"]
    if os.environ.get("THERMIQ_ORIGIN_FOLDER_ID"):
        doc_record["origin_folder_id"] = os.environ["THERMIQ_ORIGIN_FOLDER_ID"]
    db.collection("documents").document(doc_id).set(doc_record)
    print(f"  Firestore documents/{doc_id} written.")

    elapsed = time.time() - start_time
    print("\n--- Ingestion Summary ---")
    print(f"Chunks created: {len(chunks)}")
    print(f"Qdrant upserts: {upsert_count}")
    print(f"Time taken: {elapsed:.2f}s")


if __name__ == "__main__":
    main()
